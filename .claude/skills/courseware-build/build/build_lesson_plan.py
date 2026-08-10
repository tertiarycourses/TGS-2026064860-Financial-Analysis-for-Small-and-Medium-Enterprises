#!/usr/bin/env python3
"""Generate the Financial Analysis for SMEs Lesson Plan (LP) DOCX in the
Tertiary house format: cover page + Document Version Control Record + auto TOC +
Arial 11pt body + colour-coded 2-day schedule tables (9:00am-6:00pm, 8 training
hours/day, 1h lunch, tea within, final assessment Day 2 4:00pm). Slide numbers
are read from slide_index.json written by build_slides.py so the LP always
matches the current deck.
"""
import os, sys, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3
ACT=DOMAIN1+DOMAIN2+DOMAIN3
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"activities")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

# slide numbers from the current deck
IDX_PATH=os.path.join(HERE,"slide_index.json")
IDX=json.load(open(IDX_PATH)) if os.path.exists(IDX_PATH) else {}
def sl(a,b=None):
    """format a slide range from index keys"""
    if not IDX: return ""
    lo=IDX.get(a); hi=(IDX.get(b)-1) if (b and IDX.get(b)) else None
    if lo is None: return ""
    return f" (Slides {lo}–{hi})" if hi and hi>lo else f" (Slide {lo})"

BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
HEADER_FILL="1F6FEB"; TOPIC_FILL="E8F0FE"; BREAK_FILL="FFF4E5"; LUNCH_FILL="FDE9D9"; ASSESS_FILL="E8F7EE"

def act_title(n):
    a=[x for x in ACT if x["num"]==n][0]
    return f"Activity {n}: {a['title']}"

# ------------------------------------------------ schedule (single source of truth for timing)
SCHEDULE = {
 1: (C.DAY_THEMES[1], [
    ("9:00","9:30",30,"admin","Welcome, digital attendance (AM), trainer and learner introductions, ground rules, download of course material from the LMS"+sl("admin","topic1")),
    ("9:30","10:30",60,"topic","Topic 1 — Understanding Financial Statements: significance of finance, the financial system, accounting objectives and types, accounting cycle, chart of accounts, single vs double entry"+sl("topic1")),
    ("10:30","10:45",15,"break","Tea break"),
    ("10:45","12:15",90,"topic","Topic 1 (continued) — Balance sheet: assets, liabilities, equity, depreciation methods, working capital; Income statement: revenue to net income, EBITDA, cash vs accrual; Cash flow statement: operating, investing, financing"),
    ("12:15","13:00",45,"activity","Hands-on: "+act_title(1)+sl("activity1")),
    ("13:00","14:00",60,"lunch","Lunch break"),
    ("14:00","15:30",90,"topic","Digital attendance (PM). Topic 2 — Analysing Financial Ratios: ratio analysis, liquidity ratios, leverage ratios and their interpretation"+sl("topic2")),
    ("15:30","15:45",15,"break","Tea break"),
    ("15:45","16:45",60,"topic","Topic 2 (continued) — Efficiency ratios, profitability ratios, EPS, statement of changes in equity, summary of all ratios"),
    ("16:45","17:30",45,"activity","Hands-on: "+act_title(2)+sl("activity2")),
    ("17:30","18:00",30,"activity","Hands-on: "+act_title(3)+sl("activity3")+". Day 1 recap and Q&A"),
 ]),
 2: (C.DAY_THEMES[2], [
    ("9:00","9:15",15,"recap","Day 1 recap and mandatory digital attendance (AM)"),
    ("9:15","10:30",75,"topic","Topic 3 — Planning & Budgeting using Financial Statements: budgets, budgeting methods, objectives, budgeting process, capital budgeting and time value of money"+sl("topic3")),
    ("10:30","10:45",15,"break","Tea break"),
    ("10:45","11:30",45,"activity","Hands-on: "+act_title(4)+sl("activity4")),
    ("11:30","12:15",45,"activity","Hands-on: "+act_title(5)+sl("activity5")),
    ("12:15","13:00",45,"topic","Topic 3 (continued) — Cash flow analysis, pro forma statements, budget vs forecast, variances, business risk"),
    ("13:00","14:00",60,"lunch","Lunch break"),
    ("14:00","15:00",60,"topic","Digital attendance (PM). Topic 3 (continued) — Financial health of a company, investment suitability, industry ratios, benchmarking, horizontal and vertical analysis"),
    ("15:00","15:45",45,"activity","Hands-on: "+act_title(6)+sl("activity6")),
    ("15:45","16:00",15,"assess","Revision and course summary"+(f" (Slides {IDX['wrapup']}–{IDX['total']})" if IDX else "")+", course feedback and TRAQOM survey; Briefing for Assessment (Slide 14)"),
    ("16:00","17:00",60,"assess","Digital attendance (Assessment). Written Assessment (WA) — Short-Answer Questions (SAQ), 1 hour, open book"),
    ("17:00","18:00",60,"assess","Practical Performance (PP) — financial-analysis tasks, 1 hour, open book. End of class"),
 ]),
}

# ------------------------------------------------ build document
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)

prodoc.add_cover_page(doc,"LESSON PLAN",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("10.0","1 June 2026","Legacy lesson plan aligned to the v10 master trainer slides.","Han Leong"),
 ("11.0","10 August 2026",
  "Full redesign: regenerated from the single-source content pipeline; slide references aligned to the v11 visual deck; six hands-on activities scheduled across the two days.",C.TRAINER),
 (C.VERSION.lstrip("v"),C.VERSION_DATE,
  "Added realistic mock-data Excel/CSV workbooks for every activity (activities/activity01-06); slide references aligned to the v12 deck.",C.TRAINER),
])
prodoc.add_toc(doc)

def H(text,level=1):
    h=doc.add_heading(text,level=level); return h

H("Course Information",1)
info=[("Course Title",C.TITLE),("WSQ Course Reference",C.COURSE_CODE),
      ("Skills Framework",f"{C.TSC_TITLE} ({C.TSC_CODE})"),
      ("Training Provider",C.ORG+"  ("+C.UEN.replace('UEN: ','UEN ')+")"),
      ("Duration","2 days · 8 training hours per day (16 hours, including 2 hours of assessment)"),
      ("Daily Timing","9:00 am – 6:00 pm (1-hour lunch; tea breaks within training time)"),
      ("Mode","Instructor-led, with hands-on financial-analysis activities in Excel"),
      ("Trainer",C.TRAINER)]
t=doc.add_table(rows=0,cols=2); t.style="Table Grid"
for k,v in info:
    c=t.add_row().cells; c[0].text=""; r=c[0].paragraphs[0].add_run(k); r.bold=True; r.font.size=Pt(10)
    prodoc._shade_cell(c[0],TOPIC_FILL)
    c[1].text=""; c[1].paragraphs[0].add_run(v).font.size=Pt(10)

H("Learning Outcomes",1)
doc.add_paragraph("On completion of this course, learners will be able to:")
for lo in C.LEARNING_OUTCOMES:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(lo).font.size=Pt(10.5)

H("Assessment",1)
for a in [C.ASSESSMENT["written"],C.ASSESSMENT["practical"],
          "Format: Open Book — course slides, Learner Guide and approved materials only.",
          "Final assessment is conducted on Day 2 from 4:00 pm.",C.ASSESSMENT["note"]]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size=Pt(10.5)

def set_cell(cell,text,bold=False,size=9.5,color=None,fill=None,align=None):
    cell.text=""; p=cell.paragraphs[0]
    if align: p.alignment=align
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.name="Arial"
    if color: r.font.color.rgb=color
    if fill: prodoc._shade_cell(cell,fill)

KIND_FILL={"topic":TOPIC_FILL,"break":BREAK_FILL,"lunch":LUNCH_FILL,"assess":ASSESS_FILL,
           "admin":"F3F5F8","recap":"F3F5F8","activity":None}

H("Course Schedule",1)
for day,(theme,rows) in SCHEDULE.items():
    H(f"Day {day} — {theme}",2)
    tbl=doc.add_table(rows=0,cols=3); tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=tbl.add_row().cells
    for i,htext in enumerate(["Time","Duration","Topic / Activity"]):
        set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
    training=0
    for start,end,mins,kind,text in rows:
        cells=tbl.add_row().cells; fill=KIND_FILL.get(kind)
        set_cell(cells[0],f"{start}–{end}",bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        set_cell(cells[1],f"{mins} min",size=9.5,fill=fill)
        set_cell(cells[2],text,bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        if kind!="lunch": training+=mins
    for row in tbl.rows:
        row.cells[0].width=Inches(1.15); row.cells[1].width=Inches(0.9); row.cells[2].width=Inches(4.75)
    p=doc.add_paragraph(); r=p.add_run(f"Total training time: {training} minutes ({training//60} hours)."); r.italic=True; r.font.size=Pt(9.5); r.font.color.rgb=GREY
    assert training==480, f"Day {day} training minutes = {training}, expected 480"

H("Activity Reference (aligned to the TSC)",1)
tt=doc.add_table(rows=0,cols=3); tt.style="Table Grid"
hdr=tt.add_row().cells
for i,htext in enumerate(["Topic","TSC Coverage","Activities"]):
    set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
for tp in C.TOPICS:
    acts=[a for a in ACT if a["topic"]==tp["num"]]
    cells=tt.add_row().cells
    set_cell(cells[0],f"Topic {tp['code']}: {tp['title']}",bold=True,size=9.5,fill=TOPIC_FILL)
    set_cell(cells[1],tp["weighting"],size=9.5,fill=TOPIC_FILL)
    set_cell(cells[2],"; ".join(f"Activity {a['num']}: {a['title']}" for a in acts),size=9.5)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
OUT=os.path.join(REPO,"courseware",f"LP-{C.SHORT_TITLE}.docx")
doc.save(OUT)
print("Saved",OUT)
