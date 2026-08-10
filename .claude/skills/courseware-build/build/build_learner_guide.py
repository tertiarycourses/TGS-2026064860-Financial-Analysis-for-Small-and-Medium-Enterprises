#!/usr/bin/env python3
"""Generate the Financial Analysis for SMEs Learner Guide as BOTH a Markdown
mirror (LG-*.md at repo root) and a DOCX (courseware/LG-*.docx) from one source.

House format: cover page, Document Version Control Record, auto TOC, Arial 11pt
body, one section per activity (Objective · Goal · What you'll produce · Data ·
Step-by-step with workings · Test it), plus concepts per topic and a glossary.
Also writes one activity-NN-*.md file per activity into activities/.
"""
import os, sys, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3
ACT=DOMAIN1+DOMAIN2+DOMAIN3
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"activities")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

# ---------------- block DSL (single content stream → MD + DOCX) ----------------
B=[]
def h1(t): B.append(("h1",t))
def h2(t): B.append(("h2",t))
def h3(t): B.append(("h3",t))
def p(t):  B.append(("p",t))
def bullets(xs): B.append(("bullets",xs))
def steps(xs): B.append(("steps",xs))
def note(t): B.append(("note",t))
def rule(): B.append(("rule",))

# ---------------- content ----------------
h1("Introduction")
p(f"This Learner Guide accompanies the WSQ course {C.TITLE} ({C.COURSE_CODE}), conducted by {C.ORG}. "
  "It provides the key concepts for each of the three topics and detailed step-by-step instructions "
  "for all six hands-on activities. Each activity uses a worksheet downloadable from the course LMS "
  "(https://lms-tms.tertiaryinfotech.com) and is completed in Microsoft Excel or Google Sheets.")
p("Use this guide alongside the course slides. The final assessment is open book: you may refer to the "
  "slides, this Learner Guide and any approved materials, so keep your completed activity worksheets — "
  "they are your best revision notes.")

h1("Course Learning Outcomes")
bullets(C.LEARNING_OUTCOMES)

h1("Skills Framework (TSC)")
p(f"TSC Title: {C.TSC_TITLE} · TSC Code: {C.TSC_CODE}")
bullets([f"{k}: {v}" for k,v in C.TSC_KNOWLEDGE]+[f"{k}: {v}" for k,v in C.TSC_ABILITIES])

h1("Before You Start — Setup")
h3("What you need")
bullets([
 "A laptop with Microsoft Excel (2016 or later) or a Google account for Google Sheets.",
 "Access to the course LMS at https://lms-tms.tertiaryinfotech.com — log in with your registered email (an OTP is sent to you).",
 "The activity worksheets, downloaded from your course page on the LMS.",
 "A calculator (or the spreadsheet itself) for the ratio and discounting computations.",
])
h3("Conventions used in every activity")
bullets([
 "Figures in the Topic 1 activity are in $ million; Topic 2 and 3 activities use dollars.",
 "Workings shown in a box are the expected calculation or Excel formula for that step.",
 "Each activity ends with a 'Test it' check — compare your result against it before moving on.",
 "Model answers are discussed in class after each activity.",
])

# ---------------- per-topic, per-activity ----------------
for t in C.TOPICS:
    h1(f"Topic {t['code']} — {t['title']}  ({t['weighting']})")
    p(t["subtitle"])
    h3("Key concepts")
    bullets([f"{a} — {b}" for a,b in t["concepts"]])
    for a in [x for x in ACT if x["topic"]==t["num"]]:
        h2(f"Activity {a['num']} — {a['title']}")
        p(f"Maps to: {a['objective']}.")
        p(f"Goal: {a['desc']}")
        h3("What you'll produce")
        p(a["build"]+f"   (Tools: {a['services']}.)")
        h3("Data provided")
        bullets([f"{name}: {val}" for name,val in a.get("data",[])])
        h3("Step-by-step")
        steps(list(a["steps"]))
        h3("Test it")
        p(a["test"])
        note(f"The full worksheet for this activity is available on the LMS, and a printable copy is in "
             f"activities/activity-{a['num']:02d} of the course repository.")
        rule()

h1("Revision Pointers for the Final Assessment")
bullets([
 "K1/K2 — Be able to explain the statement of financial position: assets, liabilities, equity, and how the balance sheet equation stays in balance.",
 "K3 — Be able to trace the income statement from revenue to net income, and explain why profit is not cash (link to the cash flow statement).",
 "K4 — Know the statement of changes in equity: Beginning Equity + Net Income − Dividends ± Other changes = Ending Equity.",
 "K5 — Be able to pick the right analysis technique: ratio analysis, horizontal (trend) analysis, vertical analysis, benchmarking.",
 "A1 — Practise reading ratio trends across periods: what do a falling current ratio and a rising debt-to-equity together tell you?",
 "A2 — Practise the DuPont decomposition: ROE = Net Profit Margin × Asset Turnover × Equity Multiplier.",
 "Bring your completed activity worksheets — the assessment is open book.",
])

h1("Glossary")
gl=[
 ("Balance sheet","Statement of financial position: Assets = Liabilities + Equity, as at a date."),
 ("Income statement","Profit & Loss statement: Profits = Revenues − Expenses, over a period."),
 ("Cash flow statement","Cash movements in operating, investing and financing activities."),
 ("Statement of changes in equity","Reconciles opening to closing equity for the period."),
 ("Working capital","Current assets minus current liabilities."),
 ("EBITDA","Earnings before interest, tax, depreciation and amortization."),
 ("EBIT","Earnings before interest and tax — operating profit."),
 ("Current ratio","Current assets / current liabilities — a liquidity measure."),
 ("Quick (acid test) ratio","(Current assets − inventories) / current liabilities."),
 ("Debt-to-equity ratio","Total liabilities (or long-term debt) / shareholders' equity."),
 ("Inventory turnover","COGS / average inventories — stock efficiency."),
 ("ROE / ROA","Return on equity / return on assets — profitability of capital."),
 ("Time value of money","PV = FV / (1+i)^n — discounting future cash to today."),
 ("NPV","Net present value: PV of inflows − initial outlay; accept if > 0."),
 ("Profitability index","PV of inflows / PV of outflows; accept if > 1."),
 ("IRR","The discount rate at which NPV equals zero."),
 ("Payback period","Time for cumulative cash inflows to repay the initial outlay."),
 ("Variance","Difference between budget and actual — adverse or favourable."),
 ("Horizontal analysis","Percent change of the same line item across periods."),
 ("Vertical analysis","Each line as a percent of a base figure (sales or total assets)."),
 ("DuPont analysis","ROE = Net Profit Margin × Asset Turnover × Equity Multiplier."),
]
B.append(("dl",gl))

# ---------------- render Markdown ----------------
def _anchor(txt):
    return "".join(ch.lower() if ch.isalnum() else ("-" if ch in " -" else "") for ch in txt)

def render_md():
    out=[f"# {C.TITLE} — Learner Guide",""]
    out.append(f"**WSQ Course Code:** {C.COURSE_CODE}  |  **Conducted by:** {C.ORG} ({C.UEN.replace('UEN: ','UEN ')})  |  **Version {C.VERSION} · {C.VERSION_DATE}**")
    out.append("")
    out.append("## Contents"); out.append("")
    for kind,*rest in B:
        if kind=="h1": out.append(f"- [{rest[0]}](#{_anchor(rest[0])})")
        elif kind=="h2": out.append(f"  - [{rest[0]}](#{_anchor(rest[0])})")
    out.append("")
    for kind,*rest in B:
        if kind=="h1": out+=["",f"## {rest[0]}",""]
        elif kind=="h2": out+=["",f"### {rest[0]}",""]
        elif kind=="h3": out+=[f"**{rest[0]}**",""]
        elif kind=="p": out+=[rest[0],""]
        elif kind=="bullets": out+=[f"- {x}" for x in rest[0]]+[""]
        elif kind=="steps":
            for i,(instr,work) in enumerate(rest[0],1):
                out.append(f"{i}. {instr}")
                if work: out+=["",f"   ```",f"   {work}","   ```",""]
            out.append("")
        elif kind=="note": out+=[f"> **Note:** {rest[0]}",""]
        elif kind=="rule": out+=["---",""]
        elif kind=="dl":
            for term,defn in rest[0]: out.append(f"- **{term}** — {defn}")
            out.append("")
    return "\n".join(out)

MD_OUT=os.path.join(REPO,f"LG-{C.SHORT_TITLE}.md")
with open(MD_OUT,"w") as f: f.write(render_md())
print("Saved",MD_OUT)

# ---------------- per-activity markdown files in activities/ ----------------
ACT_DIR=os.path.join(REPO,"activities")
os.makedirs(ACT_DIR,exist_ok=True)
def slug(t):
    s=re.sub(r"[^a-z0-9]+","-",t.lower()).strip("-")
    return s
index_lines=[f"# Activities — {C.TITLE} ({C.COURSE_CODE})","",
             "One file per hands-on activity. The same steps, with full workings, are in the Learner Guide.",""]
for a in ACT:
    fn=f"activity-{a['num']:02d}-{slug(a['title'])}.md"
    lines=[f"# Activity {a['num']} — {a['title']}","",
           f"**Topic {a['topic']}** · {a['objective']}","",
           f"**Goal:** {a['desc']}","",
           f"**You'll produce:** {a['build']}","",
           f"**Tools:** {a['services']}","",
           "## Data provided",""]
    for name,val in a.get("data",[]):
        lines.append(f"- **{name}:** {val}")
    lines+=["","## Step-by-step",""]
    for i,(instr,work) in enumerate(a["steps"],1):
        lines.append(f"{i}. {instr}")
        if work: lines+=["","   ```",f"   {work}","   ```",""]
    lines+=["","## Test it","",a["test"],""]
    with open(os.path.join(ACT_DIR,fn),"w") as f: f.write("\n".join(lines))
    index_lines.append(f"- [Activity {a['num']} — {a['title']}]({fn})")
with open(os.path.join(ACT_DIR,"README.md"),"w") as f: f.write("\n".join(index_lines)+"\n")
print("Saved",len(ACT),"activity files +",os.path.join(ACT_DIR,"README.md"))

# ---------------- render DOCX ----------------
BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
INKCODE=RGBColor(0x0B,0x30,0x60)
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc,"LEARNER GUIDE",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("10.0","1 June 2026","Legacy learner guide aligned to the v10 master trainer slides.","Han Leong"),
 (C.VERSION.lstrip("v"),C.VERSION_DATE,
  "Full redesign: regenerated from the single-source content pipeline with six step-by-step hands-on activities (cash flow statement, ratio analysis, trend analysis, payback, NPV/PI, solvency), revision pointers and a glossary.",C.TRAINER),
])
prodoc.add_toc(doc)

def code_para(text):
    para=doc.add_paragraph()
    r=para.add_run(text); r.font.name="Consolas"; r.font.size=Pt(9.5); r.font.color.rgb=INKCODE

for kind,*rest in B:
    if kind=="h1": doc.add_heading(rest[0],level=1)
    elif kind=="h2": doc.add_heading(rest[0],level=2)
    elif kind=="h3":
        para=doc.add_paragraph(); r=para.add_run(rest[0]); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=BRAND
    elif kind=="p": doc.add_paragraph(rest[0])
    elif kind=="bullets":
        for x in rest[0]: doc.add_paragraph(x,style="List Bullet")
    elif kind=="steps":
        for i,(instr,work) in enumerate(rest[0],1):
            para=doc.add_paragraph(style="List Number"); para.add_run(instr)
            if work: code_para("Working:  "+work)
    elif kind=="note":
        para=doc.add_paragraph(); r=para.add_run("Note: "); r.bold=True; r.font.color.rgb=BRAND
        para.add_run(rest[0]).font.size=Pt(10)
    elif kind=="rule": doc.add_paragraph("")
    elif kind=="dl":
        for term,defn in rest[0]:
            para=doc.add_paragraph(style="List Bullet")
            r=para.add_run(term+" — "); r.bold=True; para.add_run(defn)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT=os.path.join(REPO,"courseware",f"LG-{C.SHORT_TITLE}.docx")
doc.save(DOCX_OUT)
print("Saved",DOCX_OUT)
