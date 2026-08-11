#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build the WSQ assessment set for 'Application Integration with Docker and Kubernetes' (TGS-2021010366):
  - Written Assessment (SAQ)  — 5 open-ended KNOWLEDGE questions (K1–K5), aligned to the slides
  - Practical Performance (PP) — 4 PRACTICAL tasks (LO1–LO4), aligned to the in-class activities
Each instrument is produced as a Question Paper and a matching Answer Key (4 DOCX total),
all with the WSQ house cover page (same as the Lesson Plan / Learner Guide). Page 1 is the cover;
page 2 carries Trainee Information + Instructions + Grading; the questions/tasks begin on page 3.
Body: Arial 11.
"""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# This script lives in the wsq-assessment skill (.claude/skills/wsq-assessment/) and runs in
# place — it detects the course repo root by walking up to the nearest dir that has a .git
# folder (or both courseware/ and assessment/). Override with env REPO=/path if needed.
def _find_repo():
    env = os.environ.get("REPO")
    if env and os.path.isdir(env):
        return os.path.abspath(env)
    d = os.path.dirname(os.path.abspath(__file__))
    while d != os.path.dirname(d):
        if os.path.isdir(os.path.join(d, ".git")) or \
           (os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "assessment"))):
            return d
        d = os.path.dirname(d)
    return os.getcwd()

REPO = _find_repo()
# prodoc.py (WSQ cover page + version control + page numbers, same as LP/LG) ships with the
# tertiary-lesson-plan skill. Look for it at the project level first, then the user level.
for _cand in (os.path.join(REPO, ".claude/skills/tertiary-lesson-plan"),
              os.path.expanduser("~/.claude/skills/tertiary-lesson-plan")):
    if os.path.exists(os.path.join(_cand, "prodoc.py")):
        sys.path.insert(0, _cand); break
import prodoc  # cover page + version control + page numbers (same as LP/LG)

# ─── EDIT PER COURSE ────────────────────────────────────────────────────────
TITLE       = "Financial Analysis for Small and Medium Enterprises"
COURSE_CODE = "TGS-2026064860"
# ────────────────────────────────────────────────────────────────────────────
# The cover page renders prodoc's module-level TGS constant. Override it so the
# assessment cover shows THIS course's ref (works with either prodoc version —
# the older project prodoc has no course_code kwarg).
prodoc.TGS = f"TGS Ref No: {COURSE_CODE}"
OUT   = os.path.join(REPO, "assessment")

# Logos: prefer the course's own courseware/assets, else fall back to the copies bundled
# in this skill (so the assessment builds even outside this project). Replace the course
# logo per course; the Tertiary Infotech logo is the same for every WSQ course.
def _logo(name):
    here = os.path.dirname(os.path.abspath(__file__))
    for p in (os.path.join(REPO, "courseware/assets", name), os.path.join(here, "assets", name)):
        if os.path.exists(p):
            return p
    return None
ORG_LOGO    = _logo("tertiary-infotech-logo.png")
COURSE_LOGO = None   # Tertiary-only cover (as LP/LG)

Q_VER, A_VER = "v6", "v6"   # v5 + activity-folder references
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)
# Assessments carry the cover page only — no Document Version Control Record.

# ---------------------------------------------------------------- WRITTEN (KNOWLEDGE)
# (criterion, context, question, [model-answer points]) — each traces to the course slides.
# Mirrors the original v3 WA: 5 open-ended SAQs mapped one-to-one to K1..K5, 1 hour, open book.
WRITTEN = [
 ("K1",
  "A small retail business is considering expanding its operations by opening a new store location. "
  "Before committing, the owner wants to evaluate the financial position of the company.",
  "As a financial analyst, how would you assess the adequacy of the company's current financial position "
  "using its statement of financial position to ensure the expansion is financially viable?",
  ["Analyse the three building blocks of the statement of financial position: assets, liabilities and equity, "
   "and confirm the accounting equation Assets = Liabilities + Equity.",
   "Assess the liquidity of the current assets (cash and equivalents, receivables, inventory) against the "
   "current liabilities — is there enough working capital to fund the expansion?",
   "Assess the sufficiency of equity (share capital and retained earnings) to support the expansion, and the "
   "level of current and long-term liabilities already carried.",
   "Compare the financial position across periods to identify trends or financial risks associated with the "
   "expansion, remembering the balance sheet is a static, book-value snapshot. "
   "(Slides: Balance Sheet — Definitions / Components / Limitations, Topic 1)"]),
 ("K2",
  "A medium-sized manufacturing company is experiencing a rapid increase in orders but is facing challenges "
  "in managing its working capital. You have been asked to review the company's balance sheet to evaluate "
  "its current financial structure.",
  "How would you interpret the relationship between assets, liabilities, and equity to determine if the "
  "company has the financial flexibility to sustain this growth?",
  ["Assess the proportion of current and non-current assets relative to current and non-current liabilities — "
   "working capital = current assets − current liabilities.",
   "Evaluate the liquidity ratios: current ratio (CA/CL) and quick ratio ((CA − inventory)/CL) to judge the "
   "ability to meet short-term obligations as orders grow.",
   "Use the debt-to-equity ratio to understand the reliance on debt versus equity financing, and whether the "
   "company has capacity to take on additional debt if necessary.",
   "A balanced structure (adequate equity, manageable leverage, positive working capital) indicates the "
   "financial flexibility to sustain growth. "
   "(Slides: The Accounting Equation, Balance Sheet Components, Liquidity Ratios, Leverage Ratios)"]),
 ("K3",
  "A small enterprise has been profitable but is facing cash flow issues. You have been tasked with "
  "analyzing the income and cash flow statements.",
  "How would you evaluate the company's profitability and cash flow to determine the root cause of the cash "
  "flow problems despite reported profits?",
  ["Compare net income from the income statement with the cash flow from operating activities — profit is "
   "accrual-based, cash flow is not.",
   "Look for discrepancies between profit and cash: high or growing receivables, excessive inventory, or "
   "delayed payments to suppliers all absorb cash while leaving profit intact.",
   "Analyse the investing and financing blocks of the cash flow statement for outflows (asset purchases, loan "
   "repayments, dividends) that drain liquidity.",
   "Review the timing of cash inflows and outflows relative to profit recognition (accrual vs cash basis). "
   "(Slides: Income Statement — Structure, Cash vs Accrual Basis, Cash Flow Statement; Activity 1)"]),
 ("K4",
  "An SME is planning to raise additional equity capital to fund a new project. Before doing so, the "
  "management wants to understand the potential impact on the company's equity structure.",
  "How would you evaluate the statement of changes in equity to advise on the effects of new equity issuance "
  "on existing shareholders and overall financial stability?",
  ["Analyse the past changes in equity: opening balance, issuance of new shares, net income/loss for each "
   "period, dividends paid and other adjustments (revaluation reserve, prior-period corrections).",
   "Consider the dilution effect of the new issue on existing shareholders' ownership percentage.",
   "Assess the impact on return on equity (ROE) — a larger equity base requires proportionally higher "
   "earnings to sustain the same return.",
   "Check the trend of equity growth and whether the company has been efficiently using its equity to "
   "generate profits (Beginning Equity + Net Income − Dividends ± Other changes = Ending Equity). "
   "(Slides: Statement of Changes in Equity, Return on Equity, Topic 2)"]),
 ("K5",
  "A startup is seeking additional investment and has presented its financial statements to potential "
  "investors. You are required to perform a financial statement analysis to evaluate the startup's "
  "financial health.",
  "How would you use various financial statement analysis techniques to provide a comprehensive evaluation "
  "that will aid investors in making informed decisions?",
  ["Apply horizontal (trend) analysis — the percentage change of each line item across periods — and vertical "
   "analysis — each line as a percentage of a base figure (sales / total assets).",
   "Use ratio analysis across the four families: liquidity (current, quick), profitability (margins, ROE, "
   "ROA), solvency/leverage (debt-to-equity, interest coverage) and efficiency (turnover ratios).",
   "Benchmark the ratios against industry standards and competitors to give the figures context.",
   "Consider non-financial factors that affect the statements: market conditions, management effectiveness "
   "and business-model viability. "
   "(Slides: Financial Statement Analysis — Methods, Benchmarking, Summary of All Ratios, Topic 3)"]),
]

# ---------------------------------------------------------------- PRACTICAL (ACTIVITY-BASED)
# Mirrors the original v3 PP: one scenario + 2 tasks mapped to A1 and A2, 1 hour, open book.
SCENARIO = (
 "You are a senior financial analyst at a consultancy firm that specializes in providing financial advice to "
 "small and medium enterprises (SMEs). One of your clients, a successful small business owner, is considering "
 "expanding their portfolio by investing in one of two potential companies: Company A (Kestrel Precision "
 "Engineering Pte Ltd) and Company B (Lighthouse Marine Supplies Pte Ltd). Their balance sheets and income "
 "statements for the last three years are provided in the data workbook 'PP Data - Financial Analysis for "
 "Small and Medium Enterprises - Company A and B.xlsx' issued with this paper. The client has asked you to "
 "perform a detailed financial analysis to identify trends and evaluate the companies' performance over "
 "multiple time periods to make an informed investment decision.")

BOX_CAP = "Paste a screenshot of your analysis in the box below"
PRACTICAL = [
 ("Task 1", "A1",
  "Using the three-year balance sheets and income statements of Company A and Company B in the data "
  "workbook provided, analyze the financial ratios or indicators over this period — for "
  "example Return on Equity (ROE), Return on Assets (ROA), Net Profit Margin, Current Ratio, and "
  "Debt-to-Equity Ratio, or others. Identify trends in these ratios and discuss how they reflect the "
  "financial health and risk associated with each company. Based on these trends, recommend which company "
  "your client should consider investing in and justify your decision. "
  "(Mirrors Activity 2 — Ratio Analysis of Two Companies, and Activity 3 — Trend Analysis; see "
  "activities/activity02 and activities/activity03.)",
  BOX_CAP,
  "Procedure (as practised in Activities 2 and 3):\n"
  "1. Lay out both companies' statements side by side in Excel, one column per year.\n"
  "2. Compute the profitability ratios per year for each company:\n"
  "   Net Profit Margin = Net Income / Sales;  ROE = Net Income / Equity;  ROA = Net Income / Total Assets.\n"
  "3. Compute the liquidity and solvency ratios per year:\n"
  "   Current Ratio = Current Assets / Current Liabilities;  Debt-to-Equity = Long-Term Debt / Equity.\n"
  "4. Compute the year-on-year growth of each ratio: (this year - last year) / last year.\n"
  "5. Identify the trends: improving margins/returns and stable liquidity favour investment; a rising\n"
  "   debt-to-equity or falling current ratio signals rising risk.\n"
  "6. Recommend the company with the stronger and more sustainable trend and justify with the ratios.\n"
  "   (Worked in class: compare ROE, ROA and net profit margin for both companies; the better\n"
  "   performer on all required metrics is the better investment.)"),
 ("Task 2", "A2",
  "Using Company A's Year 3 financial statements in the data workbook provided, calculate the DuPont analysis to break down Return "
  "on Equity (ROE) into its components: Net Profit Margin, Asset Turnover, and Equity Multiplier. Interpret "
  "each component to assess the company's operating efficiency, asset utilization, and financial leverage. "
  "Based on your analysis, explain whether Company A's financial position and performance indicate a strong "
  "potential for future growth. Provide a recommendation to your client on whether to proceed with the "
  "investment, with a focus on how the DuPont analysis supports your conclusion. "
  "(DuPont index — Topic 3, Financial Analysis for Investment Suitability; builds on Activity 2, "
  "activities/activity02.)",
  BOX_CAP,
  "DuPont decomposition: ROE = Net Profit Margin x Asset Turnover x Equity Multiplier.\n"
  "  - Operating efficiency  -> Net Profit Margin = Net Income / Sales (net income per dollar of sales).\n"
  "  - Asset utilization     -> Total Asset Turnover = Sales / Total Assets (sales per dollar of assets).\n"
  "  - Financial leverage    -> Equity Multiplier = Total Assets / Total Equity.\n"
  "Worked example (Company A, Year 3 of the data workbook: sales 140,000; net income 36,000;\n"
  "total assets 310,000; equity 135,000):\n"
  "  Net Profit Margin   = $36,000 / $140,000  = 25.71%\n"
  "  Total Asset Turnover = $140,000 / $310,000 = 45.17%\n"
  "  Equity Multiplier    = $310,000 / $135,000 = 2.30\n"
  "  ROE = 25.71% x 45.17% x 2.30 = 26.71%\n"
  "Interpretation: a healthy margin and moderate leverage drive the ROE; asset turnover shows how hard the\n"
  "asset base works. Strengths/weaknesses of the underlying metrics should be noted (net income includes\n"
  "non-cash items; total assets alone do not show productive use; total equity is not a gauge of ongoing\n"
  "profitability). Recommendation follows from whether the ROE is driven by genuine operating efficiency\n"
  "rather than leverage alone."),
]

# ---------------------------------------------------------------- doc helpers
def base_doc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    return doc

def para(doc, text, size=11, bold=False, italic=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if align is not None: p.alignment = align
    return p

def heading(doc, text, size=13):
    para(doc, text, size=size, bold=True, color=BRAND, after=6, before=8)

def answer_box(doc, lines=None, code=None, height_pt=90):
    """1x1 bordered box. `lines` → bullet-style model answer; `code` → monospace
    code/YAML/command block (indentation preserved); neither → empty answer space."""
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    if code:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in code.split("\n"):
            b = cell.add_paragraph(style=None)
            b.paragraph_format.space_after = Pt(0); b.paragraph_format.space_before = Pt(0)
            rr = b.add_run(ln if ln else " ")
            rr.font.name = "Consolas"; rr.font.size = Pt(9)
            rr._element.rPr.rFonts.set(qn('w:cs'), "Consolas")
            wt = rr._element.find(qn('w:t'))
            if wt is not None: wt.set(qn('xml:space'), 'preserve')
    elif lines:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in lines:
            b = cell.add_paragraph(style=None); b.paragraph_format.left_indent = Inches(0.15)
            rr = b.add_run("•  " + ln); rr.font.size = Pt(10.5)
    else:
        # empty answer space
        tr = t.rows[0]._tr
        trPr = tr.get_or_add_trPr(); trh = OxmlElement('w:trHeight')
        trh.set(qn('w:val'), str(int(height_pt*20))); trh.set(qn('w:hRule'), 'atLeast'); trPr.append(trh)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

FILL_GAP = 6    # extra space below each fill-in line (paired with double line spacing for writing room)

def candidate_block(doc):
    heading(doc, "Trainee Information")
    for label in ["Trainee Name (as per NRIC): ______________________________________",
                  "Last 3 digits and alphabet of NRIC/FIN: ____________________",
                  "Date: ____________________"]:
        p = para(doc, label, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

# Assessment briefing (from the course slides — "Briefing for Assessment").
BRIEFING = [
    "Place phones and other materials under the table or on the floor.",
    "No photos or recording of assessment scripts.",
    "No discussion during the assessment.",
    "Use a black/blue pen for hard-copy assessments.",
    "No liquid paper / correction tape.",
    "Scripts are collected when time is up.",
]

LMS_URL = "https://lms-tms.tertiaryinfotech.com/"

def add_hyperlink(p, url, text):
    """Add a real clickable Word hyperlink (blue, underlined) to paragraph p."""
    r_id = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)  # 11pt
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run); p._p.append(link)
    return link

def instructions(doc, minutes_text):
    heading(doc, "Instructions to Candidate")
    # None marks the upload instruction, which carries a clickable LMS hyperlink.
    items = [
        "This is an individual exercise.",
        "This is an open-book assessment.",
        f"A total of {minutes_text} is given to complete this assessment.",
        None,
    ] + BRIEFING
    for i, s in enumerate(items, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        if s is None:
            p.add_run(f"{i}.  Complete your answers on the document provided and "
                      "upload the completed answers to the LMS at ").font.size = Pt(11)
            add_hyperlink(p, LMS_URL, LMS_URL)
            p.add_run(".").font.size = Pt(11)
        else:
            p.add_run(f"{i}.  {s}").font.size = Pt(11)

def grading(doc, what):
    heading(doc, "Grading")
    para(doc, what, size=11, after=12)
    for ln in ["Grade: _______  (C / NYC)",
               "Assessor Name: __________________________   Assessor NRIC: ________________",
               "Date: ________________________                    Signature: ____________________"]:
        p = para(doc, ln, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

def finish(doc, path):
    prodoc.add_page_numbers(doc); prodoc.enable_update_fields(doc)
    doc.save(path); print("  saved:", os.path.basename(path))

# ---------------------------------------------------------------- builders
def build_wa(answers):
    doc = base_doc()
    kind = "Written Assessment (SAQ) — Answer Key" if answers else "Written Assessment (SAQ)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Written Assessment (SAQ)" if answers else "Written Assessment (SAQ)",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; questions begin on the next page.
        candidate_block(doc); instructions(doc, "1 hour")
        grading(doc, "Candidate has answered all written questions and demonstrated the underpinning "
                     "knowledge required for the course learning outcomes.")
        page_break(doc)
    para(doc, "Short-Answer Questions (Knowledge)", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Answer all questions in your own words. Each question tests underpinning knowledge covered in the "
              "course slides.", size=10.5, italic=True, color=GREY, after=8)
    # Pagination is EXPLICIT — two questions to a page on the paper, one model answer to a
    # page in the key. Do not swap this for Word's keepNext/cantSplit: Word pushes an
    # oversized box to the next page, but Google Docs draws the border anyway and prints the
    # question text and the page footer straight THROUGH it. See SKILL.md → Pagination.
    per_page = 1 if answers else 2
    for i, (crit, ctx, q, pts) in enumerate(WRITTEN, 1):
        para(doc, f"Question {i}:", size=11.5, bold=True, after=2, before=6)
        para(doc, ctx, size=11, after=3)
        para(doc, f"{q}  ({crit})", size=11, bold=True, after=4)
        answer_box(doc, lines=pts if answers else None)
        if i % per_page == 0 and i < len(WRITTEN):
            page_break(doc)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to WA (SAQ) - {TITLE} - {suffix}.docx" if answers
            else f"WA (SAQ) - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

def build_pp(answers):
    doc = base_doc()
    kind = "Practical Performance (PP) — Answer Key" if answers else "Practical Performance (PP)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Practical Performance Assessment" if answers else "Practical Performance Assessment",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; the problem begins on the next page.
        candidate_block(doc); instructions(doc, "1 hour")
        grading(doc, "Candidate has successfully completed all PP tasks and can explain the overall "
                     "functions and features used to achieve them.")
        page_break(doc)
    para(doc, "Practical Problem", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Scenario", size=11.5, bold=True, after=2)
    para(doc, SCENARIO, size=11, after=8)
    # Practical tasks are long and their boxes are tall, so they get a page each — on the
    # paper AND in the key. Same rule as the WA: the page break is ours, not the renderer's.
    for i, (label, crit, prompt, cap, pts) in enumerate(PRACTICAL, 1):
        para(doc, f"{label} ({crit}):", size=11.5, bold=True, after=2, before=6)
        para(doc, prompt, size=11, after=3)
        para(doc, cap, size=10.5, italic=True, color=GREY, after=4)
        answer_box(doc, code=pts if answers else None, height_pt=150)
        if i < len(PRACTICAL):
            page_break(doc)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to PP Assessment - {TITLE} - {suffix}.docx" if answers
            else f"PP Assessment - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

if __name__ == "__main__":
    print("Building WSQ assessment set…")
    build_wa(answers=False); build_wa(answers=True)
    build_pp(answers=False); build_pp(answers=True)
    print(f"Done. WA: {len(WRITTEN)} questions · PP: {len(PRACTICAL)} tasks.")
